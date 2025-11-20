import re              # Standard module for Regular Expressions (used for pattern matching/search)
import csv             # Standard module for reading/writing CSV files
import os              # Standard module for interacting with the operating system (paths, files, etc.)
import glob            # Standard module for finding files matching a pattern (like '*.docx')
import sys             # Standard module for system-specific parameters and functions (used for exit, executable path)
import subprocess      # Standard module for running new applications/commands (used to run 'pip install')

# --- Auto-Installation Function ---
def install_module(package):
    """Installs a single required module using pip."""
    print(f"\n--- Dependency Missing: Installing '{package}' ---")
    try:
        # Use sys.executable to ensure pip runs within the active venv/environment
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
        print(f"--- SUCCESS: Successfully installed {package}. ---")
    except subprocess.CalledProcessError as e:
        # Handles installation failure (e.g., no internet, permissions)
        print(f"--- CRITICAL ERROR: Failed to install {package}. Please check your internet connection or permissions. ---")
        print(f"Error details: {e}")
        sys.exit(1) # Exit the script upon critical failure

# --- Dependency Check and Initial Imports ---
try:
    # Attempt to import the primary required module from the python-docx package
    from docx import Document
except ImportError:
    # If the docx module is not found, install the required package
    install_module('python-docx')
    # Try the import again after installation
    try:
        from docx import Document
    except ImportError:
        # Catch-all for scenarios where installation succeeded but import still fails 
        # (e.g., due to a local file named docx.py conflicting with the library)
        print("\n--- CRITICAL ERROR: The 'Document' class could not be imported.")
        print("Please check if you have a file named 'docx.py' in this directory and rename it. ---")
        sys.exit(1)


# --- REMAINDER OF THE ORIGINAL SCRIPT (Client ID and Case Note Logic) ---

# Master list of all keys (used to define the boundaries of the data sections in the DOCX file)
STOPPER_KEYS = [
    'Client', 'HMIS #', 'Entry Date', 'Exit Date', 'Room #', 'Contact', 'DOB', 'Email',
    'Case Manager', 'Required Information', 'Benefits', 'Food Stamps',
    'Income', 'SSI', 'SSDI', 'Trimet Tickets', 'Health Insurance',
    'Summary:', 'Tier 1 Warnings', 'Case Notes', 'Intake', 'HMIS Privacy'
]
# Creates a massive Regular Expression pattern to search for any of the STOPPER_KEYS
LOOKAHEAD_PATTERN = '|'.join([re.escape(k) for k in STOPPER_KEYS])


def extract_client_identifiers(docx_filepath):
    """
    Extracts Client Name and ID (HMIS #) only from the filename, 
    assuming the format: 'Client Name #ID.docx'.
    """
    header_data = {}
    # Get the filename without the path and extension
    filename_base = os.path.splitext(os.path.basename(docx_filepath))[0]
    # Regex to find a string followed by an optional space, a '#', and then digits
    name_id_match = re.match(r'(.+?)\s*#(\d+)', filename_base)
    # Extract Client Name (group 1) or use the whole filename if no match
    header_data['Client Name'] = name_id_match.group(1).strip() if name_id_match else filename_base
    # Extract Client ID (group 2) or set to 'N/A'
    header_data['Client ID'] = name_id_match.group(2) if name_id_match else 'N/A'
    return header_data


# --- CASE NOTE EXTRACTION (Retained and robust) ---

def _extract_raw_notes(docx_filepath):
    """
    Reads a DOCX file and extracts raw case notes based on bold formatting,
    starting after 'Case Notes:' is found.
    """
    try:
        # Load the document object from the file
        document = Document(docx_filepath)
    except Exception as e:
        # Handles issues with opening the DOCX file
        print(f"Error opening DOCX file for notes extraction: {e}")
        return []

    all_notes = []
    current_note = None # Stores the note currently being assembled
    in_case_notes_section = False # Flag to start processing only after 'Case Notes:' is seen

    # Regex to find a date pattern (e.g., 1/1/2025 or 01/01/25)
    date_pattern = re.compile(r'(\d{1,2}/+\d{1,2}/+\d{2,4})')

    def process_paragraph_list(paragraphs, current_note, all_notes, in_case_notes_section, date_pattern):
        # Nested function to handle paragraphs (from main body or inside tables)
        for para in paragraphs:
            # Look for the start of the case notes section
            if not in_case_notes_section and 'Case Notes:' in para.text:
                in_case_notes_section = True
                continue # Skip the "Case Notes:" header paragraph itself

            if not in_case_notes_section:
                continue # Skip all paragraphs before the 'Case Notes:' header

            full_paragraph_text = para.text.strip()
            if not full_paragraph_text:
                continue # Skip empty paragraphs

            is_note_start = False
            bold_text_accumulator = ""

            # Check for bold text at the start of the paragraph
            for run in para.runs:
                if run.bold:
                    bold_text_accumulator += run.text
                else:
                    # Stop if non-bold text is encountered (end of header)
                    break 

            bold_text_cleaned = bold_text_accumulator.strip()
            # Check if the accumulated bold text contains a date pattern
            date_match = date_pattern.search(bold_text_cleaned)

            # A new note starts if a date is found AND the bold text starts with that date
            if date_match and bold_text_cleaned.startswith(date_match.group(0)):
                is_note_start = True

                if current_note:
                    # Clean up and finalize the *previous* note before starting a new one
                    current_note['Note'] = re.sub(r'\s+', ' ', current_note['Note']).strip()
                    all_notes.append(current_note)

                # Extract and clean the Date
                date = re.sub(r'/+', '/', date_match.group(1)).strip()
                # Extract the Staff name (text after the date)
                staff_raw = bold_text_cleaned[date_match.end():].strip()
                # Remove trailing colons, spaces, or slashes from staff name
                staff = re.sub(r'[:\s/]+$', '', staff_raw)

                try:
                    # Find the content of the note (text following the bold header)
                    header_end_index = para.text.index(bold_text_accumulator) + len(bold_text_accumulator)
                    note_content = para.text[header_end_index:].strip()
                except ValueError:
                    # Fallback: if bold text not found (shouldn't happen here), use full text
                    note_content = full_paragraph_text 

                # Initialize the new current_note dictionary
                current_note = {
                    'Date': date,
                    'Staff': staff,
                    'Note': note_content
                }

            # If it's not a new note start, but we are inside a note, append the paragraph text
            elif not is_note_start and current_note:
                current_note['Note'] += " " + full_paragraph_text
        
        # Return the updated state
        return current_note, all_notes, in_case_notes_section

    # Process paragraphs inside tables first (as the Case Notes might be in a table)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # Recursively process the paragraphs within each cell
                current_note, all_notes, in_case_notes_section = process_paragraph_list(
                    cell.paragraphs, current_note, all_notes, in_case_notes_section, date_pattern
                )

    # Process top-level paragraphs (those not contained in any table)
    current_note, all_notes, in_case_notes_section = process_paragraph_list(
        document.paragraphs, current_note, all_notes, in_case_notes_section, date_pattern
    )

    # Finalize the very last note that was still held in current_note
    if current_note:
        current_note['Note'] = re.sub(r'\s+', ' ', current_note['Note']).strip()
        all_notes.append(current_note)

    return all_notes


# --- BATCH EXECUTION (Simplified) ---
def run_batch_conversion(base_directory="."):
    """
    Finds all DOCX files and writes a CSV report containing only
    Client Name, HMIS #, and Case Notes for each.
    """

    # Find all .docx files recursively (case-insensitive search)
    docx_files = glob.glob(os.path.join(base_directory, '**', '*.docx'), recursive=True)
    docx_files += glob.glob(os.path.join(base_directory, '**', '*.DOCX'), recursive=True)

    if not docx_files:
        print(f"\n--- ERROR: No DOCX files found. ---")
        return

    print(f"\nFound {len(docx_files)} DOCX files to process.")

    # Loop through every found DOCX file
    for input_file in docx_files:
        print(f"--- Processing: {input_file} ---")

        base_dir = os.path.dirname(input_file) # Directory of the input file
        # Filename without extension
        base_name = os.path.splitext(os.path.basename(input_file))[0] 
        # Create output filename in the same directory
        output_file = os.path.join(base_dir, f"{base_name}_CASENOTES_extracted.csv")

        # 1. EXTRACT DATA
        header_data = extract_client_identifiers(input_file) # From filename
        raw_notes = _extract_raw_notes(input_file) # From document content

        # 2. WRITE LEAN CSV
        try:
            # Open the new CSV file for writing
            with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)

                # Write Client & HMIS # in key/value format (A1, B1, A2, B2)
                writer.writerow(['Client', header_data.get('Client Name', '')])
                writer.writerow(['HMIS #', header_data.get('Client ID', '')])

                # Write empty lines for visual separation in the CSV
                for _ in range(4):
                    writer.writerow(['', '']) 
                
                # --- Write Case Notes Section Header ---
                writer.writerow(['Case Notes', ''])
                # Write the column headers for the notes data
                writer.writerow(['Date', 'Staff', 'Note'])

                # Write each extracted note as a new row
                for note in raw_notes:
                    writer.writerow([
                        note.get('Date', ''),
                        note.get('Staff', ''),
                        note.get('Note', '')
                    ])

            print(f"--- SUCCESS: extracted casenote: {os.path.basename(output_file)}")
        except Exception as e:
            print(f"--- ERROR: Failed to write CSV for {input_file}. Error: {e} ---")

        print("-" * 20)

# Standard entry point for a Python script
if __name__ == "__main__":
    STARTING_DIRECTORY = '.' # Look for DOCX files in the current directory and subdirectories
    
    # Run the main conversion process
    run_batch_conversion(STARTING_DIRECTORY)