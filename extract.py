import re
import os
import sys
import zipfile
import glob
import shutil
from datetime import datetime, date

# Try importing the necessary library, providing instructions if it's missing
try:
    from docx import Document # The library needed to read content from Word (.docx) files
except ImportError:
    # If python-docx is not installed, print an error and stop the script
    print("\n--- CRITICAL ERROR: 'python-docx' not installed. ---")
    print("This library is required to read the contents of the DOCX files.")
    print("Please install it: pip install python-docx")
    sys.exit(1)

# --- Configuration ---
# A temporary folder name used to extract the contents of the input ZIP
TEMP_DIR = 'temp_zip_extraction_dir'

def is_case_file(docx_file_path):
    """
    Checks if a DOCX file contains "Entry Date" or "Exit Date" keywords
    to confirm it is a structured client case file.
    """

    try:
        document = Document(docx_file_path)
    except Exception:
        # If the file is corrupt or not a true DOCX structure, skip it
        return False

    # Check the main body paragraphs for the keywords
    for para in document.paragraphs:
        if any(term in para.text for term in ["Entry Date", "Exit Date"]):
            return True

    # Also check content inside tables for the keywords
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(term in cell.text for term in ["Entry Date", "Exit Date"]):
                    return True

    return False # If no date terms are found anywhere, it's not considered a case file


def filter_and_copy_to_folder(input_zip_filepath, output_folder_path):
    """
    Manages the entire process: extraction, filtering, and copying of matching files
    to the specified output folder, preserving the structure.
    """

    # 1. Prepare environment
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR) # Delete the temporary directory if it already exists
    os.makedirs(TEMP_DIR, exist_ok=True) # Create a fresh, empty temporary directory

    # 2. Prepare output folder
    if os.path.exists(output_folder_path):
        print(f"Warning: Deleting existing output folder: {output_folder_path}")
        shutil.rmtree(output_folder_path) # Delete output folder if it exists (for clean run)
    os.makedirs(output_folder_path, exist_ok=True) # Create a fresh, empty output directory

    files_found = 0
    files_matched = 0

    print(f"\n--- Starting File Filtering for: {os.path.basename(input_zip_filepath)} ---")

    try:
        # 3. Extract contents of the input zip into the temporary directory
        with zipfile.ZipFile(input_zip_filepath, 'r') as zip_ref:
            zip_ref.extractall(TEMP_DIR)

        # 4. Find all .docx files (case-insensitive) inside the temporary directory recursively
        docx_files = glob.glob(os.path.join(TEMP_DIR, '**', '*.docx'), recursive=True)
        docx_files += glob.glob(os.path.join(TEMP_DIR, '**', '*.DOCX'), recursive=True)
        files_found = len(docx_files)

        if not docx_files:
            print(f"No DOCX files found inside the ZIP file.")
            return

        # 5. Process each DOCX file
        for docx_path in docx_files:
            # Get the path of the file relative to the temp folder (this preserves the original folder structure)
            relative_path = os.path.relpath(docx_path, TEMP_DIR)

            print(f"  [Checking] {relative_path}...")

            if is_case_file(docx_path):
                # Check passed! Define the destination path in the new output folder
                destination_path = os.path.join(output_folder_path, relative_path)

                # Ensure the subdirectories in the output folder exist before copying
                os.makedirs(os.path.dirname(destination_path), exist_ok=True)

                # Copy the file from the temporary location to the final output folder
                shutil.copy2(docx_path, destination_path) # copy2 attempts to preserve metadata

                files_matched += 1
                print(f"  [MATCHED] Copied {relative_path} to output folder.")
            else:
                # Check failed. Skip this file.
                print(f"  [Skipped] {relative_path}: Missing 'Entry/Exit Date' tag.")


    except FileNotFoundError:
        print(f"\n--- ERROR: Input ZIP file not found at '{input_zip_filepath}' ---")
        return
    except Exception as e:
        print(f"\n--- FATAL ERROR during processing: {e} ---")
        return
    finally:
        # 6. Cleanup: Always delete the temporary directory and its contents
        if os.path.exists(TEMP_DIR):
            shutil.rmtree(TEMP_DIR)

    # 7. Report success and statistics
    print("\n" + "=" * 60)
    print(f"FILTERING COMPLETE!")
    print(f"Total DOCX files found in input: {files_found}")
    print(f"Total Case Note files saved: {files_matched}")
    print(f"Output Folder created: {os.path.abspath(output_folder_path)}")
    print("=" * 60)


if __name__ == "__main__":

    # --- Collect User Inputs ---
    # Check if the input file name was passed directly as a command-line argument
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        # Otherwise, ask the user for the input file name
        input_file = input("Enter the name of the input ZIP file (e.g., 'archive.zip'): ").strip()

    # Ask the user for the desired name of the new output folder
    output_folder = input("Enter the name for the NEW output FOLDER (e.g., 'casenotes_only'): ").strip()

    if not output_folder:
        # Use a default name if the user leaves the output name blank
        print("Output folder name cannot be empty. Using 'filtered_case_notes_folder'.")
        output_folder = 'filtered_case_notes_folder'

    # Run the main filtering and copying function
    filter_and_copy_to_folder(input_file, output_folder)
